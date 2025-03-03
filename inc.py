import pandas as pd
from docx import Document
import re
import os
import zipfile
from pathlib import Path
import datetime
import numpy as np

def merge_employee_data_and_zip(excel_file, word_template, output_folder, zip_name=None):
    """
    Read employee data from Excel, replace repeated placeholders in Word document,
    and create a single zip file containing all documents

    Args:
        excel_file (str): Path to Excel file with employee data
        word_template (str): Path to Word template with placeholders
        output_folder (str): Folder to save generated documents and final zip
        zip_name (str, optional): Name for the zip file. If None, uses current date
    """
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Read Excel data
    df = pd.read_excel(excel_file)

    # Define the mapping between Excel columns and Word placeholders
    placeholder_mapping = {
        'Emp ID': '[Employee ID]',
        'Name': '[Name]',
        'Department': '[Employee Department]',
        'Employee Title': '[Employee Title]',
        'Employee Type': '[Employee Type]',
        '2024 Bonus': '[Bonus in INR]',
        'Basic Salary': '[Basic in INR]',
        'HRA': '[HRA in INR]',
        'Other Allowences': '[Other Allowance in INR]',
        'Provident Fund': '[Provident Fund in INR]',
        'Company Deposit': '[Company Deposit in INR]',
        'Total Fixed': '[Total Fixed in INR]',
        'Bonus 2025 (At Target)': '[Target in INR]',
        'Total CTC': '[Total CTC in INR]'
    }

    # Current date in the format "February 26, 2025"
    current_date = datetime.datetime.now().strftime("%B %d, %Y")

    # Create a temp folder for documents
    docs_folder = os.path.join(output_folder, "temp_docs")
    os.makedirs(docs_folder, exist_ok=True)

    # List to keep track of all generated files
    generated_files = []

    # Process each row in the Excel file
    for index, row in df.iterrows():
        # Create a copy of the template for each employee
        doc = Document(word_template)

        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            # Replace date placeholder
            if '[Date]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[Date]', current_date)

            # Replace unique employee data placeholders
            for excel_col, word_placeholder in placeholder_mapping.items():
                if word_placeholder in paragraph.text:
                    # Get the value from Excel
                    if excel_col in df.columns:
                        value = row[excel_col]
                        # Check if the value is empty or NaN
                        if pd.notna(value) and value != "":
                            formatted_value = str(value)
                            # Replace the placeholder with the formatted value
                            paragraph.text = paragraph.text.replace(word_placeholder, formatted_value)

        # Also check tables for placeholders
        for table in doc.tables:
            for row_idx in range(len(table.rows)):
                for col_idx in range(len(table.columns)):
                    cell = table.cell(row_idx, col_idx)

                    # Replace date placeholder
                    if '[Date]' in cell.text:
                        cell.text = cell.text.replace('[Date]', current_date)

                    # Replace unique employee data placeholders
                    for excel_col, word_placeholder in placeholder_mapping.items():
                        if word_placeholder in cell.text:
                            # Get the value from Excel
                            if excel_col in df.columns:
                                value = row[excel_col]
                                # Check if value is empty or NaN
                                if pd.notna(value) and value != "":
                                    formatted_value = str(value)
                                    # Replace the placeholder with the formatted value
                                    cell.text = cell.text.replace(word_placeholder, formatted_value)

        # Get employee ID for naming files
        emp_id = str(row['Emp ID'])
        safe_emp_id = re.sub(r'[^\w\s-]', '', emp_id).strip().replace(' ', '_')

        # Save the document with employee ID as name
        doc_path = os.path.join(docs_folder, f"{safe_emp_id}.docx")
        doc.save(doc_path)
        generated_files.append(doc_path)

        print(f"Generated document for employee ID: {emp_id}")

    # Create a single zip file containing all documents
    if zip_name is None:
        today = datetime.datetime.now().strftime("%Y%m%d")
        zip_name = f"employee_documents_{today}.zip"

    zip_path = os.path.join(output_folder, zip_name)
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in generated_files:
            # Add file to zip with just the filename (not the full path)
            filename = os.path.basename(file)
            zipf.write(file, arcname=filename)

    print(f"Processed {len(df)} employees. All documents saved in single zip file: {zip_path}")

    # Optionally remove the temporary folder with individual files
    import shutil
    shutil.rmtree(docs_folder)

    return zip_path

# Example usage
if __name__ == "__main__":
    excel_file = "employee_data.xlsx"  # Path to your Excel file
    word_template = "template.docx"  # Path to your Word template
    output_folder = "output"  # Folder to save the final zip

    # Create the documents and zip file
    zip_file = merge_employee_data_and_zip(excel_file, word_template, output_folder)
    print(f"ZIP file created at: {zip_file}")
